"""
Word document generation module.
Creates formatted Word documents with original and translated text.
"""

import logging
from pathlib import Path
from typing import List, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# Arabic text processing
import arabic_reshaper
from bidi.algorithm import get_display
from reportlab.lib import colors

# Import UI configuration
from ui_config import ui_config

logger = logging.getLogger(__name__)


# Custom SimpleDocTemplate to disable page numbers
class SimpleDocTemplateNoPageNumbers(SimpleDocTemplate):
    def afterPageSetup(self):
        # Override this method to prevent page numbering
        pass

class WordDocumentGenerator:
    """Generates Word documents with original and translated text."""

    def __init__(self):
        self._arabic_fonts_registered = False
        self._registered_font = 'Helvetica'

    def _setup_arabic_fonts(self):
        """Register Arabic-compatible fonts for PDF generation."""
        if self._arabic_fonts_registered:
            return

        try:
            # Try to register Noto Sans Arabic first (best Arabic support)
            font_paths = [
                '/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf',
                '/usr/share/fonts/truetype/noto/NotoSansArabic-Bold.ttf',
                '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
            ]

            registered_font = None
            for font_path in font_paths:
                try:
                    if 'NotoSansArabic' in font_path:
                        font_name = 'NotoSansArabic-Bold' if 'Bold' in font_path else 'NotoSansArabic'
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        registered_font = 'NotoSansArabic'
                        addMapping('NotoSansArabic', 0, 0, 'NotoSansArabic')
                        addMapping('NotoSansArabic', 1, 0, 'NotoSansArabic-Bold')
                    elif 'Liberation' in font_path:
                        font_name = 'LiberationSans-Bold' if 'Bold' in font_path else 'LiberationSans'
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        registered_font = 'LiberationSans'
                        addMapping('LiberationSans', 0, 0, 'LiberationSans')
                        addMapping('LiberationSans', 1, 0, 'LiberationSans-Bold')
                    elif 'DejaVu' in font_path:
                        font_name = 'DejaVuSans-Bold' if 'Bold' in font_path else 'DejaVuSans'
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        registered_font = 'DejaVuSans'
                        addMapping('DejaVuSans', 0, 0, 'DejaVuSans')
                        addMapping('DejaVuSans', 1, 0, 'DejaVuSans-Bold')

                    if registered_font:
                        break

                except Exception as font_error:
                    continue

            self._registered_font = registered_font or 'Helvetica'
            self._arabic_fonts_registered = True
            logger.info(f"Arabic fonts registered successfully: {self._registered_font}")

        except Exception as e:
            logger.warning(f"Failed to register Arabic fonts: {e}")
            self._registered_font = 'Helvetica'
            self._arabic_fonts_registered = True

    async def create_clean_arabic_document(
        self,
        translated_pairs: List[Tuple[str, str]],
        output_path: Path,
        original_filename: str = "document"
    ) -> Path:
        """
        Create a simple Word document with only Arabic translations - plain text format.

        Args:
            translated_pairs: List of (original_text, translated_text) tuples
            output_path: Path where to save the document
            original_filename: Name of the original file for reference

        Returns:
            Path to the created document

        Raises:
            Exception: If document creation fails
        """
        try:
            # Create a new document
            doc = Document()

            # Process each translation - ONLY Arabic text, very simple
            for i, (original_text, translated_text) in enumerate(translated_pairs, 1):
                # Clean Arabic translation (remove numbers and formatting)
                clean_arabic = self._clean_arabic_translation(translated_text)

                # Add ONLY Arabic translation - simple paragraph
                arabic_para = doc.add_paragraph(clean_arabic)
                arabic_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                arabic_para.runs[0].font.size = Pt(18)

            # Save the document
            doc.save(str(output_path))

            logger.info(f"Created simple Arabic-only document with {len(translated_pairs)} translations")
            return output_path

        except Exception as e:
            logger.error(f"Failed to create Arabic Word document: {e}")
            raise Exception(f"Arabic document generation failed: {str(e)}")

    async def create_bilingual_document(
        self,
        translated_pairs: List[Tuple[str, str]],
        output_path: Path,
        original_filename: str = "document"
    ) -> Path:
        """
        Create a Word document with original English and Arabic translations in proper format.

        Args:
            translated_pairs: List of (original_text, translated_text) tuples
            output_path: Path where to save the document
            original_filename: Name of the original file for reference

        Returns:
            Path to the created document

        Raises:
            Exception: If document creation fails
        """
        try:
            # Create a new document
            doc = Document()

            # Add header information at the top
            header_info = doc.add_paragraph()
            header_run1 = header_info.add_run(f"Translation Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            header_run1.font.size = Pt(12)
            header_run1.font.name = 'Times New Roman'
            
            header_run2 = header_info.add_run("Developer: @dextermorgenk")
            header_run2.font.size = Pt(12)
            header_run2.font.name = 'Times New Roman'
            header_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            doc.add_paragraph()

            # Process each translation pair with improved formatting
            for i, (original_text, translated_text) in enumerate(translated_pairs, 1):
                await self._add_structured_translation_pair(doc, i, original_text, translated_text)

            # Save the document
            doc.save(str(output_path))

            logger.info(f"Created structured bilingual document with {len(translated_pairs)} translation pairs")
            return output_path

        except Exception as e:
            logger.error(f"Failed to create Word document: {e}")
            raise Exception(f"Document generation failed: {str(e)}")

    async def _add_structured_translation_pair(
        self,
        doc,
        index: int,
        original_text: str,
        translated_text: str
    ):
        """Add a structured translation pair maintaining original document format."""
        try:
            # Clean translated text first
            clean_arabic = self._clean_arabic_translation(translated_text)

            # Check if this is a title/heading (typically shorter text or contains certain formatting)
            is_heading = len(original_text.strip()) < 100 and (
                original_text.isupper() or 
                original_text.strip().endswith(':') or
                any(word in original_text.lower() for word in ['chapter', 'section', 'part', 'lab', 'experiment'])
            )

            if is_heading:
                # Format as heading
                heading_para = doc.add_paragraph()
                
                # English heading
                eng_run = heading_para.add_run(original_text)
                eng_run.font.name = 'Times New Roman'
                eng_run.font.size = Pt(ui_config.current_arabic_font_size + 2)
                eng_run.font.bold = True
                eng_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Add line break
                heading_para.add_run('\n')
                
                # Arabic heading
                ar_run = heading_para.add_run(clean_arabic)
                ar_run.font.name = 'Arial Unicode MS'
                ar_run.font.size = Pt(ui_config.current_arabic_font_size + 2)
                ar_run.font.bold = True
                ar_run.font.color.rgb = RGBColor(0, 0, 0)
                
                heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            else:
                # Format as regular content with proper structure
                content_para = doc.add_paragraph()
                
                # English text
                eng_run = content_para.add_run(original_text)
                eng_run.font.name = 'Times New Roman'
                eng_run.font.size = Pt(ui_config.current_arabic_font_size)
                eng_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Add spacing between English and Arabic
                content_para.add_run('\n')
                
                # Arabic text
                ar_run = content_para.add_run(clean_arabic)
                ar_run.font.name = 'Arial Unicode MS'
                ar_run.font.size = Pt(ui_config.current_arabic_font_size)
                ar_run.font.color.rgb = RGBColor(0, 0, 0)
                
                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Add spacing between sections
            if index < len(translated_pairs):
                spacing_para = doc.add_paragraph()
                spacing_para.add_run(' ')
                spacing_para.space_after = Pt(6)

        except Exception as e:
            logger.warning(f"Failed to add structured translation pair {index}: {e}")
            # Fallback to simple format
            await self._add_translation_pair(doc, index, original_text, translated_text)

    async def _add_translation_pair(
        self,
        doc,
        index: int,
        original_text: str,
        translated_text: str
    ):
        """Add a translation pair to the document with proper paragraph formatting."""
        try:
            # Add English text paragraph (preserving original structure)
            english_para = doc.add_paragraph()
            english_run = english_para.add_run(original_text)
            english_run.font.name = 'Times New Roman'
            english_run.font.size = Pt(ui_config.current_arabic_font_size)
            english_run.font.color.rgb = RGBColor(0, 0, 0)
            english_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Clean translated text (preserve math symbols and structure)
            clean_arabic = self._clean_arabic_translation(translated_text)

            # Add Arabic translation paragraph directly below
            arabic_para = doc.add_paragraph()
            arabic_run = arabic_para.add_run(clean_arabic)
            arabic_run.font.name = 'Arial Unicode MS'
            arabic_run.font.size = Pt(ui_config.current_arabic_font_size)
            arabic_run.font.color.rgb = RGBColor(0, 0, 0)
            arabic_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right align for Arabic

            # Add spacing between translation blocks
            doc.add_paragraph()

        except Exception as e:
            logger.warning(f"Failed to add translation pair {index}: {e}")
            # Add basic clean text as fallback
            clean_arabic = self._clean_arabic_translation(translated_text)
            fallback_para = doc.add_paragraph(f"{original_text}")
            fallback_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fallback_para2 = doc.add_paragraph(f"{clean_arabic}")
            fallback_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _set_rtl_alignment(self, paragraph):
        """Set right-to-left alignment for Arabic text."""
        try:
            # Set paragraph alignment to right
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Set RTL direction
            pPr = paragraph._element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

        except Exception as e:
            logger.warning(f"Failed to set RTL alignment: {e}")
            # Fallback to right alignment only
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _clean_arabic_translation(self, arabic_text: str) -> str:
        """Clean Arabic translation while preserving math expressions and symbols."""
        import re

        # Only remove leading numbering added by translation (like "1-", "2.", "3)" at the beginning)
        text = re.sub(r'^\s*\d+[-.)]\s*', '', arabic_text)

        # Remove any brackets with numbers added by translation [1], [2], etc.
        text = re.sub(r'\[\d+\]', '', text)

        # Remove any remaining leading/trailing whitespace
        text = text.strip()

        # Preserve mathematical expressions and symbols
        # Keep equations, formulas, variables, and mathematical notation intact
        # This includes: equations (x=y), functions (sin, cos), symbols (∑, ∫, etc.)
        
        return text

    def create_error_document(self, error_message: str, output_path: Path) -> Path:
        """Create a document with error information."""
        try:
            doc = Document()

            title = doc.add_heading('Translation Error', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            error_para = doc.add_paragraph()
            error_para.add_run("An error occurred during translation:\n\n").bold = True
            error_para.add_run(error_message)

            timestamp_para = doc.add_paragraph()
            timestamp_para.add_run(f"\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(str(output_path))
            return output_path

        except Exception as e:
            logger.error(f"Failed to create error document: {e}")
            raise

    async def create_arabic_pdf(
        self,
        translated_pairs: List[Tuple[str, str]],
        output_path: Path,
        original_filename: str = "document"
    ) -> Path:
        """
        Create a very simple PDF document with only Arabic text - plain text only.

        Args:
            translated_pairs: List of (original_text, translated_text) tuples
            output_path: Path where to save the PDF
            original_filename: Name of the original file for reference

        Returns:
            Path to the created PDF document

        Raises:
            Exception: If PDF creation fails
        """
        try:
            # Create PDF document with simple margins and no page numbers
            doc = SimpleDocTemplateNoPageNumbers(str(output_path), pagesize=A4,
                                                 rightMargin=50, leftMargin=50,
                                                 topMargin=50, bottomMargin=50)

            # Create styles
            styles = getSampleStyleSheet()

            # Register Arabic-supporting fonts
            self._setup_arabic_fonts()

            # Very simple text style using UI config
            arabic_style = ParagraphStyle(
                'PlainArabicText',
                parent=styles['Normal'],
                alignment=TA_RIGHT,
                fontSize=ui_config.PDF_STYLES['arabic_size'],
                spaceAfter=ui_config.PDF_STYLES['spacing'],
                spaceBefore=0,
                fontName=self._registered_font,
                leading=ui_config.PDF_STYLES['arabic_size'] + 6,
                textColor=colors.black,
                borderWidth=0,
                borderColor=None,
                backColor=None
            )

            # Build plain content - Arabic text only
            content = []

            # Header information style - centered English text
            header_style = ParagraphStyle(
                'HeaderInfo',
                parent=styles['Normal'],
                alignment=TA_CENTER,
                fontSize=14,  # Increased font size
                spaceAfter=3,
                spaceBefore=0,
                fontName=self._registered_font,
                leading=18,  # Increased leading
                textColor=colors.black
            )

            # Add header information in English (centered at top) - no filename
            content.append(Paragraph(f"Translation Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", header_style))
            content.append(Paragraph("Developer: @dextermorgenk", header_style))
            content.append(Spacer(1, 30)) # Increased spacing

            # Add only Arabic translations - completely plain
            for i, (original_text, translated_text) in enumerate(translated_pairs, 1):
                # Clean Arabic translation
                clean_arabic = self._clean_arabic_translation(translated_text)
                display_arabic = self._clean_text_for_pdf(clean_arabic)

                # Add simple Arabic text paragraph
                content.append(Paragraph(display_arabic, arabic_style))

                # Small space between paragraphs
                content.append(Spacer(1, 10)) # Increased spacing

            # Build PDF - completely plain
            doc.build(content)

            logger.info(f"Created plain Arabic PDF with {len(translated_pairs)} translations")
            return output_path

        except Exception as e:
            logger.error(f"Failed to create Arabic PDF: {e}")
            raise Exception(f"Arabic PDF generation failed: {str(e)}")

    def _clean_text_for_pdf(self, text: str) -> str:
        """Clean and properly shape Arabic text for PDF rendering."""
        # Remove problematic characters
        text = text.replace('\u200f', '')  # Remove RTL mark
        text = text.replace('\u200e', '')  # Remove LTR mark
        text = text.replace('\ufeff', '')  # Remove BOM

        # Handle mathematical and scientific notations by ensuring they are rendered correctly
        # This might involve specific regex or library calls if complex parsing is needed
        # For now, we focus on common cases and rely on the Arabic reshaping for general text.
        # Example: Ensure LaTeX-like math notation is preserved as much as possible
        # This is a basic approach; advanced math rendering might require dedicated libraries.

        try:
            # Check if text contains Arabic characters
            if any('\u0600' <= char <= '\u06FF' for char in text):
                # Properly reshape Arabic text for correct display
                reshaped_text = arabic_reshaper.reshape(text)
                # Apply bidirectional algorithm for proper RTL display
                display_text = get_display(reshaped_text)
                return display_text
        except Exception as e:
            logger.warning(f"Failed to reshape Arabic text: {e}")
            # If reshaping fails, return original text
            pass

        return text.strip()


    async def create_bilingual_pdf(
        self,
        translated_pairs: List[Tuple[str, str]],
        output_path: Path,
        original_filename: str = "document"
    ) -> Path:
        """
        Create a structured bilingual PDF document matching original format.

        Args:
            translated_pairs: List of (original_text, translated_text) tuples
            output_path: Path where to save the PDF
            original_filename: Name of the original file for reference

        Returns:
            Path to the created PDF document
        """
        try:
            # Create PDF document with proper margins
            doc = SimpleDocTemplateNoPageNumbers(str(output_path), pagesize=A4,
                                                 rightMargin=72, leftMargin=72,
                                                 topMargin=72, bottomMargin=72)

            # Create styles
            styles = getSampleStyleSheet()

            # Register Arabic-supporting fonts
            self._setup_arabic_fonts()

            # Header style
            header_style = ParagraphStyle(
                'HeaderInfo',
                parent=styles['Normal'],
                alignment=TA_CENTER,
                fontSize=12,
                spaceAfter=5,
                spaceBefore=0,
                fontName=self._registered_font,
                leading=16,
                textColor=colors.black
            )

            # Heading style for titles/sections
            heading_style = ParagraphStyle(
                'SectionHeading',
                parent=styles['Heading2'],
                alignment=TA_CENTER,
                fontSize=16,
                spaceAfter=12,
                spaceBefore=12,
                fontName=self._registered_font,
                textColor=colors.black,
                borderWidth=0
            )

            # English content style
            english_style = ParagraphStyle(
                'EnglishContent',
                parent=styles['Normal'],
                alignment=TA_JUSTIFY,
                fontSize=ui_config.PDF_STYLES['english_size'],
                spaceAfter=4,
                spaceBefore=0,
                fontName=self._registered_font,
                leading=ui_config.PDF_STYLES['english_size'] + 4,
                textColor=colors.black
            )

            # Arabic content style
            arabic_style = ParagraphStyle(
                'ArabicContent',
                parent=styles['Normal'],
                alignment=TA_RIGHT,
                fontSize=ui_config.PDF_STYLES['arabic_size'],
                spaceAfter=8,
                spaceBefore=4,
                fontName=self._registered_font,
                leading=ui_config.PDF_STYLES['arabic_size'] + 6,
                textColor=colors.black
            )

            # Build structured content
            content = []

            # Add header information
            content.append(Paragraph(f"Translation Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", header_style))
            content.append(Paragraph("Developer: @dextermorgenk", header_style))
            content.append(Spacer(1, 24))

            # Process each translation pair with structure detection
            for i, (original_text, translated_text) in enumerate(translated_pairs, 1):
                clean_arabic = self._clean_arabic_translation(translated_text)
                display_arabic = self._clean_text_for_pdf(clean_arabic)

                # Detect if this is a heading/title
                is_heading = len(original_text.strip()) < 100 and (
                    original_text.isupper() or 
                    original_text.strip().endswith(':') or
                    any(word in original_text.lower() for word in ['chapter', 'section', 'part', 'lab', 'experiment'])
                )

                if is_heading:
                    # Format as heading
                    combined_heading = f"{original_text}<br/>{display_arabic}"
                    content.append(Paragraph(combined_heading, heading_style))
                    content.append(Spacer(1, 12))
                else:
                    # Format as regular content
                    if i > 1:
                        content.append(Spacer(1, 6))
                    
                    content.append(Paragraph(original_text, english_style))
                    content.append(Paragraph(display_arabic, arabic_style))

            # Build PDF
            doc.build(content)

            logger.info(f"Created structured bilingual PDF with {len(translated_pairs)} translation pairs")
            return output_path

        except Exception as e:
            logger.error(f"Failed to create bilingual PDF: {e}")
            raise Exception(f"Structured bilingual PDF generation failed: {str(e)}")

    def _is_arabic(self, text: str) -> bool:
        """Check if the text contains Arabic characters."""
        return any('\u0600' <= char <= '\u06FF' for char in text)