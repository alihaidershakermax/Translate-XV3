"""
File handling module for PDF and Word document processing.
Handles text extraction from various file formats.
"""

import logging
import tempfile
from pathlib import Path
from typing import Optional, List
import io

# PDF processing
import pdfplumber

# Word document processing
from docx import Document

logger = logging.getLogger(__name__)


class FileProcessor:
    """Handles file processing operations for PDF and Word documents."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx']
    
    def _is_page_number(self, text: str) -> bool:
        """Check if text is likely a page number."""
        import re
        
        # Remove whitespace
        text = text.strip()
        
        # Check for common page number patterns
        patterns = [
            r'^\d+$',                    # Just a number (1, 2, 3)
            r'^Page\s+\d+$',            # "Page 1", "Page 2"
            r'^\d+\s*/\s*\d+$',         # "1/10", "2 / 15"
            r'^\-\s*\d+\s*\-$',         # "- 5 -", "-10-"
            r'^\[\s*\d+\s*\]$',         # "[1]", "[ 5 ]"
            r'^\(\s*\d+\s*\)$',         # "(1)", "( 5 )"
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
                
        return False
    
    async def extract_text_from_file(self, file_path: Path) -> List[str]:
        """
        Extract text from a file and return as list of lines.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            List of text lines extracted from the file
            
        Raises:
            ValueError: If file format is not supported
            Exception: If text extraction fails
        """
        try:
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                return await self._extract_text_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return await self._extract_text_from_word(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise
    
    async def _extract_text_from_pdf(self, file_path: Path) -> List[str]:
        """Extract text from PDF file."""
        try:
            lines = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        # Split by lines and filter empty lines and page numbers
                        page_lines = []
                        for line in text.split('\n'):
                            line = line.strip()
                            # Skip empty lines and standalone page numbers
                            if line and not self._is_page_number(line):
                                page_lines.append(line)
                        lines.extend(page_lines)
            
            if not lines:
                raise ValueError("No text content found in PDF")
                
            logger.info(f"Extracted {len(lines)} lines from PDF")
            return lines
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_text_from_word(self, file_path: Path) -> List[str]:
        """Extract text from Word document."""
        try:
            lines = []
            
            # Load the document
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    lines.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            lines.append(text)
            
            if not lines:
                raise ValueError("No text content found in Word document")
            
            logger.info(f"Extracted {len(lines)} lines from Word document")
            return lines
            
        except Exception as e:
            logger.error(f"Word document text extraction failed: {e}")
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def validate_file(self, file_path: Path, max_size: int) -> bool:
        """
        Validate file format and size.
        
        Args:
            file_path: Path to the file
            max_size: Maximum file size in bytes
            
        Returns:
            True if file is valid, False otherwise
        """
        try:
            # Check if file exists
            if not file_path.exists():
                logger.error(f"File does not exist: {file_path}")
                return False
            
            # Check file size
            file_size = file_path.stat().st_size
            if file_size > max_size:
                logger.error(f"File too large: {file_size} bytes (max: {max_size})")
                return False
            
            # Check file format
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_extension}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"File validation failed: {e}")
            return False
